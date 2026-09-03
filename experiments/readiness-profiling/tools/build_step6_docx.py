from __future__ import annotations

import csv, json, math, sys
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]
OUT=Path(sys.argv[1]) if len(sys.argv)>1 else ROOT/'step6/report/ANFA_Step_6_Capacity_Actuation_Delay_Detailed_Report.docx'
CHARTS=ROOT/'step6/docx-charts'; CHARTS.mkdir(parents=True,exist_ok=True)
BLUE=RGBColor(46,116,181); DARK=RGBColor(31,77,120); NAVY=RGBColor(11,37,69); GRAY=RGBColor(90,99,112); GREEN=RGBColor(22,101,52); AMBER=RGBColor(122,90,0)
NS='http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def font(run,size=11,bold=None,color=None,italic=None):
    run.font.name='Calibri'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color is not None: run.font.color.rgb=color
    if italic is not None: run.italic=italic

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_margins(cell,top=80,start=120,bottom=80,end=120):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for side,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        e=tcMar.find(qn('w:'+side))
        if e is None: e=OxmlElement('w:'+side); tcMar.append(e)
        e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa')

def geometry(table,widths):
    table.autofit=False; tblPr=table._tbl.tblPr
    layout=tblPr.find(qn('w:tblLayout'))
    if layout is None: layout=OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'),'fixed')
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    ind=tblPr.find(qn('w:tblInd'))
    if ind is None: ind=OxmlElement('w:tblInd'); tblPr.append(ind)
    ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for w in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcW=cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tcW is None: tcW=OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'),str(widths[i])); tcW.set(qn('w:type'),'dxa'); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc,headers,rows,widths,size=8.8):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); shade(c,'F2F4F7')
        for r in c.paragraphs[0].runs: font(r,size,bold=True,color=NAVY)
    trPr=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement('w:tblHeader'); repeat.set(qn('w:val'),'true'); trPr.append(repeat)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: font(r,size)
    geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(0); return t

def heading(doc,text,level=1):
    p=doc.add_paragraph(text,style=f'Heading {level}'); return p

def para(doc,text,bold_prefix=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.1
    if bold_prefix and text.startswith(bold_prefix): font(p.add_run(bold_prefix),11,bold=True,color=NAVY); text=text[len(bold_prefix):]
    font(p.add_run(text),11); return p

def callout(doc,label,text,color=GREEN):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.08); p.paragraph_format.right_indent=Inches(.08); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F4F6F9'); pPr.append(shd)
    font(p.add_run(label+': '),10.5,bold=True,color=color); font(p.add_run(text),10.5)

def page_field(p):
    f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),'PAGE'); p._p.append(f)

def make_chart(path,title,categories,series,ymax=None,reference=None,ylabel='Seconds'):
    W,H=1300,720; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    f=ImageFont.truetype('arial.ttf',22); sm=ImageFont.truetype('arial.ttf',18); tf=ImageFont.truetype('arialbd.ttf',32)
    L,R,T,B=105,50,110,105; vals=[v for arr in series.values() for v in arr]; top=ymax or max(vals)*1.18
    px=lambda i:L+i*(W-L-R)/max(1,len(categories)-1); py=lambda v:T+(top-v)*(H-T-B)/top
    d.text((W/2,30),title,font=tf,fill='#1F4D78',anchor='ma')
    for i in range(6):
        v=top*i/5; y=py(v); d.line((L,y,W-R,y),fill='#E5E7EB',width=2); d.text((L-10,y),f'{v:.1f}',font=sm,fill='#374151',anchor='rm')
    d.line((L,T,L,H-B),fill='#111827',width=3); d.line((L,H-B,W-R,H-B),fill='#111827',width=3)
    colors=['#166534','#DC2626','#2563EB','#7C3AED']
    for (name,arr),color in zip(series.items(),colors):
        pts=[(px(i),py(v)) for i,v in enumerate(arr)]; d.line(pts,fill=color,width=5)
        for x,y in pts:d.ellipse((x-6,y-6,x+6,y+6),fill=color)
    if reference is not None:
        y=py(reference); d.line((L,y,W-R,y),fill='#B45309',width=3); d.text((W-R,y-8),f'Horizon {reference:g}s',font=sm,fill='#B45309',anchor='rs')
    for i,c in enumerate(categories): d.text((px(i),H-B+20),str(c),font=sm,fill='#374151',anchor='ma')
    x=L
    for (name,_),color in zip(series.items(),colors): d.line((x,85,x+32,85),fill=color,width=5); d.text((x+42,85),name,font=sm,fill='#374151',anchor='lm'); x+=270
    d.text((25,H/2),ylabel,font=f,fill='#374151',anchor='mm'); im.save(path)

cached={int(r['increment']):r for r in csv.DictReader(open(ROOT/'step6/report/cached/delay-distributions.csv',encoding='utf-8')) if r['metric']=='trial_effective_serving_delay_s'}
ready={int(r['increment']):r for r in csv.DictReader(open(ROOT/'step6/report/cached/delay-distributions.csv',encoding='utf-8')) if r['metric']=='trial_readiness_delay_s'}
repull={r['metric']:r for r in csv.DictReader(open(ROOT/'step6/report/registry-repull/delay-distributions.csv',encoding='utf-8')) if r['increment']=='1'}
make_chart(CHARTS/'increment_delays.png','Actuation delay by replica increment',['+1','+2','+3'],{'P95 Ready':[float(ready[i]['p95']) for i in (1,2,3)],'P95 Effective service':[float(cached[i]['p95']) for i in (1,2,3)]},ymax=9,reference=9)
make_chart(CHARTS/'cache_compare.png','Pre-pulled versus registry repull',['Median Ready','P95 Ready','Median Service','P95 Service'],{'Pre-pulled':[1.506,2.291,2.755,3.248],'Registry repull':[2.403,2.513,3.895,4.553]},ymax=5.5)

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,color,bef,aft in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color; s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT; font(header.add_run('ANFA Research | Step 6 Capacity Actuation Delay'),9,color=GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(footer.add_run('Page '),9,color=GRAY); page_field(footer)

for _ in range(6): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('CAPACITY ACTUATION\nDELAY REPORT'),29,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('ANFA Research Project - Step 6'),15,color=DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Measured Kubernetes readiness and first-service delay'),11,italic=True,color=GRAY)
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Local kind / Docker Desktop campaign\nCompleted 4 August 2026'),10.5,bold=True,color=DARK)
doc.add_page_break()

heading(doc,'Executive result',1)
para(doc,'Step 6 measured how long a scaling instruction takes to become usable serving capacity. Thirty cached-image trials covered scale-ups from one Pod to two, three, and four Pods. Ten additional registry-repull trials quantified image availability effects. All retained main trials were valid.')
add_table(doc,['Decision','Final local result','Reason'],[
    ('Forecast horizon','9 seconds','ceil(6.514 s P95 effective service + 2 s minimum margin)'),
    ('Operational metric','Effective first service','Covers readiness plus actual Service use by every requested new Pod'),
    ('Image treatment','Pre-pull and verify','Registry repull increased +1-Pod P95 service by 1.305 s'),
    ('Transfer rule','Remeasure on native K3s','The local laptop topology and bridge are development-specific'),
],[1900,2300,5160],9)
callout(doc,'Final local decision','Use a 9-second forecast horizon for the local kind experiments. Replace this value with the native-K3s P95-plus-margin result before the main research campaign.')

heading(doc,'Contents',1)
add_table(doc,['Section','Subject'],[(i+1,x) for i,x in enumerate(['Purpose and relationship to earlier steps','Frozen configuration','Timestamp authority and delay definitions','Experimental methodology','Instrumentation validation','Cached campaign results','Replica-increment comparison','Image cache comparison','Forecast-horizon calculation','Findings and interpretation','Limitations','Corrections and exclusions','Reproducibility artifacts','Completion assessment','Appendix: distribution tables'])],[900,8460],9)

heading(doc,'1. Purpose and relationship to earlier steps',1)
para(doc,'Step 4 created the deterministic benchmark application. Step 5 measured how much load each ready Pod count can safely serve. Step 6 answers the remaining timing question: how early must scaling begin so that the required Pods are actually ready and serving before forecast traffic arrives?')
callout(doc,'Connection','Step 5 determines how many Pods are needed; Step 6 determines how early those Pods must be requested.',DARK)

heading(doc,'2. Frozen configuration',1)
add_table(doc,['Item','Frozen value'],[
    ('Cluster','kind-anfa-dev on Docker Desktop / WSL2; two experiment workers'),('Application','Go benchmark service v0.1.0'),('Image','anfa/benchmark-app@sha256:0fd880c...61cee'),('Work','50,000 SHA-256 iterations; seed anfa-benchmark-v1'),('Resources','500m CPU and 128Mi memory request=limit'),('Readiness probe','HTTP /readyz every 1 second'),('Service','NodePort 30080 through documented local bridge'),('Step 5 capacity','Cpod=45 RPS; C2=90; C3=105; C4=130 RPS'),('Manifest pull policy','IfNotPresent; image pre-pulled and verified before experiments')
],[2200,7160],9)

heading(doc,'3. Timestamp authority and delay definitions',1)
add_table(doc,['Timestamp','Event','Authority'],[
    ('t_forecast','Forecast becomes available','Harness UTC + monotonic clock'),('t_decision','Controller/harness selects replica target','Harness UTC + monotonic clock'),('t_scale_sent','Deployment scale request issued','Harness UTC + monotonic clock'),('t_scale_ack','Kubernetes API acknowledges update','Harness UTC + monotonic clock'),('t_created','Pod object created','metadata.creationTimestamp'),('t_scheduled','PodScheduled=True','Condition lastTransitionTime'),('t_started','Container running','containerStatuses.state.running.startedAt'),('t_ready','Pod Ready=True','Condition lastTransitionTime'),('t_app_ready','Application enables readiness','X-Benchmark-Ready-At'),('t_first_request','New Pod first serves /work','Client monotonic clock + X-Benchmark-Pod-UID')
],[1800,3800,3760],8.4)
para(doc,'Kubernetes lifecycle fields have whole-second resolution. They are retained for authoritative ordering, while end-to-end creation, readiness, and serving delays use the harness monotonic first-observation clock. This prevents misleading negative differences when whole-second API fields and sub-second client timestamps occur in the same second.')
add_table(doc,['Delay','Definition'],[
    ('Decision','t_decision - t_forecast'),('Deployment API','t_scale_ack - t_scale_sent'),('Creation','first observation of new Pod - t_scale_sent'),('Scheduling','t_scheduled - t_created'),('Startup','t_started - t_scheduled'),('Container-to-Ready','t_ready - t_started'),('Readiness actuation','Ready first observed - t_scale_sent'),('Effective serving','first /work response from new Pod - t_scale_sent'),('Increment completion','maximum delay among all requested new Pods')
],[2600,6760],9)

heading(doc,'4. Experimental methodology',1)
para(doc,'Each cached trial established exactly one Ready baseline Pod, waited for recovery, emitted forecast and decision markers, scaled to the selected target, polled Kubernetes state, and sent bounded fresh-connection /work probes until every new Pod was Ready and had served. The Deployment was then returned to one Pod.')
add_table(doc,['Factor','Method'],[
    ('Replica increments','1→2, 1→3, and 1→4'),('Repetitions','10 valid trials per increment; 30 total'),('Ordering','Rotated 2/3/4, 3/4/2, 4/2/3 blocks'),('Polling','Nominal 100 ms; actual command and network overhead retained'),('Recovery','15 seconds for the main repeated block'),('Timeout','180 seconds'),('Completion','Last required new Pod Ready and then first observed serving'),('Evidence','Immutable trial directories with Pods, events, nodes, EndpointSlices, CSV and JSON')
],[2100,7260],9)

heading(doc,'5. Instrumentation validation',1)
para(doc,'Instrumentation pilots were deliberately excluded. They detected a stopped NodePort bridge, overlapping scale processes, HTTP connection reuse that pinned requests to one backend, and whole-second timestamp subtraction. The final runner uses a fail-fast bridge check, one isolated trial at a time, fresh bounded curl connections, Pod UID headers, and monotonic observation times.')
callout(doc,'Scientific treatment','No failed pilot was silently deleted. Invalid or superseded runs remain preserved and are disclosed separately from the 30 main trials.',AMBER)

heading(doc,'6. Cached campaign results',1)
add_table(doc,['Statistic','Overall effective serving delay'],[('n','30 trials'),('Minimum','2.387 s'),('Median','3.569 s'),('P90','5.219 s'),('P95','6.514 s'),('Maximum','6.829 s')],[3600,5760],9.5)
doc.add_picture(str(CHARTS/'increment_delays.png'),width=Inches(6.35)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.inline_shapes[-1]._inline.docPr.set('descr','Line chart comparing P95 readiness and effective-serving delay for scale increments of one, two, and three Pods against the selected nine-second horizon.')

heading(doc,'7. Replica-increment comparison',1)
add_table(doc,['Scale','Metric','Median','P90','P95','Maximum'],[
    ('1→2','Ready','1.506','2.176','2.291','2.291'),('1→2','Effective service','2.755','3.242','3.248','3.248'),('1→3','Ready','1.868','2.257','2.270','2.270'),('1→3','Effective service','3.997','5.219','5.715','5.715'),('1→4','Ready','2.249','2.807','2.816','2.816'),('1→4','Effective service','4.274','6.514','6.829','6.829')
],[1100,2100,1500,1500,1500,1660],8.5)
para(doc,'Readiness varied little across increments, but effective service increased as more Pods were added because the completion rule required every new backend to be sampled. This is operationally useful: a controller needs requested capacity to participate in traffic, not merely to have a Ready condition.')

heading(doc,'8. Image cache comparison',1)
para(doc,'All 60 new Pods in the cached campaign had matching “image already present” events and the expected image ID. A temporary local registry then supported a separate 0→1 treatment Deployment with the identical binary and runtime settings. Ten image-reference-absent registry-repull repetitions were retained; first-use rep00 was excluded.')
add_table(doc,['Treatment','n','Median Ready','P95 Ready','Median service','P95 service'],[
    ('Pre-pulled +1','10','1.506','2.291','2.755','3.248'),('Registry repull +1','10','2.403','2.513','3.895','4.553')
],[2200,700,1500,1500,1700,1760],8.7)
doc.add_picture(str(CHARTS/'cache_compare.png'),width=Inches(6.35)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.inline_shapes[-1]._inline.docPr.set('descr','Line chart comparing median and P95 readiness and effective-serving delays for pre-pulled and local-registry-repull image treatments.')
para(doc,'Registry-repull P95 effective service was 1.305 seconds slower. Pull duration median was 1.236 seconds and P95/maximum was 1.310 seconds. The excluded first-use pilot reached 7.344 seconds Ready and 8.772 seconds effective service because pulling began roughly six seconds after scheduling.')
callout(doc,'Image decision','Pre-pull and verify the immutable image on every experiment worker. Keep IfNotPresent until the native-K3s distribution method is validated; do not introduce registry delay into the main treatment.',GREEN)

heading(doc,'9. Forecast-horizon calculation',1)
para(doc,'The selected metric is trial-level effective serving delay, because it includes Kubernetes actuation plus observable Service participation by every requested new Pod. The predeclared rule is:')
callout(doc,'Formula','H = ceil(P95 effective serving delay + max(2 seconds, 20% of P95)).',DARK)
add_table(doc,['Term','Value'],[('P95 effective serving delay','6.514 s'),('20% of P95','1.303 s'),('Minimum safety margin','2.000 s'),('Selected margin','2.000 s'),('Unrounded total','8.514 s'),('Final local horizon','9 seconds')],[3400,5960],10)
para(doc,'The slowest-increment 1→4 P95 is 6.829 seconds. With ten observations, nearest-rank P95 equals that increment’s maximum; adding two seconds still rounds to the same 9-second horizon. Therefore the decision is robust to using either the pooled operational distribution or the slowest supported increment.')

heading(doc,'10. Principal findings',1)
add_table(doc,['Finding','Interpretation'],[
    ('API acknowledgement is small','P95 scale API delay remained below 0.265 s.'),('Pods become Ready quickly','Cached P95 Ready was 2.29–2.82 s across increments.'),('Ready is not equivalent to serving','Last-backend sampling raised P95 service to 3.25–6.83 s.'),('Increment size matters','More new Pods require more backend observations and increase tail delay.'),('Image availability matters','Registry repull added 1.305 s to +1-Pod P95 service.'),('Nine seconds is conservative locally','It covers measured P95 plus the declared two-second margin.')
],[2800,6560],9)

heading(doc,'11. Limitations',1)
add_table(doc,['Limitation','Consequence'],[
    ('One physical laptop','All logical nodes share CPU, memory, cooling, WSL2, and Docker Desktop.'),('Local NodePort/socat bridge','First-service path differs from native K3s.'),('100 ms requested polling','kubectl and HTTP execution make first observation an upper bound.'),('Whole-second Kubernetes fields','Sub-second scheduler/startup components cannot be estimated precisely.'),('Harness decision marker','Not yet the real predictive controller’s computation delay.'),('Ten trials per increment','Nearest-rank per-increment P95 equals the maximum.'),('Registry repull treatment','May reuse containerd layer content; not a remote WAN cold pull.'),('Post-campaign cleanup','Containerd restarts occurred after measurement and do not affect retained windows.')
],[2600,6760],8.7)

heading(doc,'12. Corrections and exclusions',1)
add_table(doc,['Event','Classification','Treatment'],[
    ('Bridge stopped','Invalid pilot','Fail-fast preflight added; excluded'),('Overlapping runners','Invalid pilot','Serial campaign enforced; excluded'),('Persistent HTTP connection','Biased pilot method','Replaced with bounded fresh curl connections'),('Negative API-derived creation value','Resolution artifact','Monotonic first-observation used for totals'),('Registry rep00','First-use pilot','Disclosed separately; excluded from ten-repetition distribution'),('Progress-array error','Automation bookkeeping','Two valid trial artifacts preserved; campaign resumed at rep03'),('Image cleanup/reload issue','Post-experiment operations','No retained measurement changed; frozen source policy restored')
],[2200,1900,5260],8.2)

heading(doc,'13. Reproducibility artifacts',1)
add_table(doc,['Artifact','Purpose'],[
    ('STEP6.md','Protocol, timestamp model, rules, and decisions'),('step6/experiment-config.json','Machine-readable campaign configuration'),('scripts/run-step6-actuation.ps1','Single cached scale-up collector'),('scripts/run-step6-campaign.ps1','Rotated 30-trial campaign driver'),('scripts/run-step6-cold-trial.ps1','Registry-repull treatment collector'),('tools/analyze_step6.py','Nearest-rank aggregation and horizon calculation'),('step6/runs/*','Immutable raw trial directories'),('step6/report/cached/*','Cached normalized datasets and distributions'),('step6/report/registry-repull/*','Repull normalized datasets and distributions'),('step6/report/step6-final-result.json','Final machine-readable decision')
],[3000,6360],8.7)

heading(doc,'14. Completion assessment',1)
add_table(doc,['Completion criterion','Status'],[
    ('Readiness delay measured, not assumed','COMPLETE'),('Lifecycle timestamps preserved','COMPLETE'),('Multiple increments repeated','COMPLETE - 10 each'),('Median/P90/P95/maximum reported','COMPLETE'),('First-service delay measured','COMPLETE'),('Image-cache behaviour evaluated','COMPLETE'),('Pre-pull decision documented','COMPLETE'),('Forecast horizon tied to evidence','COMPLETE - 9 s local'),('Native-K3s calibration','REQUIRED BEFORE MAIN CAMPAIGN')
],[6500,2860],9)

heading(doc,'Appendix A. Detailed cached delay distributions',1)
rows=[]
for r in csv.DictReader(open(ROOT/'step6/report/cached/delay-distributions.csv',encoding='utf-8')):
    rows.append((r['increment'],r['metric'],r['n'],f"{float(r['median']):.3f}",f"{float(r['p90']):.3f}",f"{float(r['p95']):.3f}",f"{float(r['maximum']):.3f}"))
add_table(doc,['+Pods','Metric','n','Median','P90','P95','Max'],rows,[700,3100,700,1200,1200,1200,1260],7.7)

heading(doc,'Appendix B. Final handoff',1)
para(doc,'For the local kind environment, forecasts must provide expected workload at least 9 seconds ahead of the arrival time. The controller converts that workload into the empirical Step 5 replica requirement, requests replicas at or before the horizon boundary, and verifies that the required Pods become Ready and serve before traffic arrives. The final native-K3s campaign must repeat Step 5 capacity profiling and Step 6 delay profiling; its values supersede the local constants.')
callout(doc,'Final local constants','Cpod=45 RPS; empirical capacities C1=45, C2=90, C3=105, C4=130 RPS; forecast horizon H=9 seconds; immutable image pre-pulled on every worker.',GREEN)

OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
