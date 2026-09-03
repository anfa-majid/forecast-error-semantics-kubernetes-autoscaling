from pathlib import Path
import argparse, csv, json, math
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
parser = argparse.ArgumentParser(description="Build the readiness-profiling report from a complete cloud evidence directory.")
parser.add_argument("--cloud-directory", type=Path, default=ROOT / "step6" / "cloud")
parser.add_argument("--output", type=Path, default=ROOT / "documentation" / "Step_6_Cloud_K3s_Capacity_Actuation_Delay_Detailed_Report.docx")
args = parser.parse_args()
CLOUD = args.cloud_directory.resolve()
RUNS = CLOUD / "runs"
OUT = args.output.resolve()
NAVY=RGBColor(11,37,69); BLUE=RGBColor(46,116,181); DARK=RGBColor(31,77,120); GRAY=RGBColor(90,99,112); GREEN=RGBColor(22,101,52); AMBER=RGBColor(146,96,0)

def percentile(values,p):
    v=sorted(values); x=(p/100)*(len(v)-1); lo=math.floor(x); hi=math.ceil(x)
    return v[lo] if lo==hi else v[lo]+(v[hi]-v[lo])*(x-lo)

progress=list(csv.DictReader(open(CLOUD/'campaign-progress.csv',encoding='utf-8')))
valid=[r for r in progress if r['status']=='valid']
for r in valid:
    for k in ('readiness_delay_s','effective_serving_delay_s','maximum_probe_gap_ms'): r[k]=float(r[k])
groups={n:[r for r in valid if int(r['target_replicas'])==n] for n in (2,3,4)}
dist=[]
for n,rows in groups.items():
    for label,key in [('Readiness','readiness_delay_s'),('Effective service','effective_serving_delay_s')]:
        vals=[r[key] for r in rows]
        dist.append((n,label,len(vals),percentile(vals,50),percentile(vals,90),percentile(vals,95),max(vals)))
slowest_p95=max(x[5] for x in dist if x[1]=='Effective service')
uncertainty=max(r['maximum_probe_gap_ms'] for r in valid)/1000
margin=max(2.0,0.2*slowest_p95); horizon=math.ceil(slowest_p95+uncertainty+margin)

pod_rows=[]
for d in RUNS.glob('cloud-cached-1to*-rep*-*'):
    p=d/'per-pod.csv'
    s=d/'trial-summary.json'
    if p.exists() and s.exists() and json.loads(s.read_text(encoding='utf-8'))['valid']:
        pod_rows.extend(list(csv.DictReader(open(p,encoding='utf-8'))))

def font(run,size=10.5,bold=None,color=None,italic=None):
    run.font.name='Calibri'; rpr=run._element.get_or_add_rPr(); rpr.rFonts.set(qn('w:ascii'),'Calibri'); rpr.rFonts.set(qn('w:hAnsi'),'Calibri'); run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color is not None: run.font.color.rgb=color
    if italic is not None: run.italic=italic

def shade(cell,fill):
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(shd)

def geometry(table,widths):
    table.autofit=False; pr=table._tbl.tblPr
    layout=OxmlElement('w:tblLayout'); layout.set(qn('w:type'),'fixed'); pr.append(layout)
    w=OxmlElement('w:tblW'); w.set(qn('w:w'),str(sum(widths))); w.set(qn('w:type'),'dxa'); pr.append(w)
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); pr.append(ind)
    grid=table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for width in widths:
        c=OxmlElement('w:gridCol'); c.set(qn('w:w'),str(width)); grid.append(c)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcpr=cell._tc.get_or_add_tcPr(); tcw=tcpr.find(qn('w:tcW')) or OxmlElement('w:tcW')
            if tcw.getparent() is None: tcpr.append(tcw)
            tcw.set(qn('w:w'),str(widths[i])); tcw.set(qn('w:type'),'dxa'); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            mar=OxmlElement('w:tcMar')
            for side,val in [('top',80),('bottom',80),('start',120),('end',120)]:
                x=OxmlElement('w:'+side); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); mar.append(x)
            tcpr.append(mar)

def table(doc,headers,rows,widths,size=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text=str(h); shade(t.rows[0].cells[i],'EAF0F7')
        for r in t.rows[0].cells[i].paragraphs[0].runs: font(r,size,bold=True,color=NAVY)
    hp=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement('w:tblHeader'); repeat.set(qn('w:val'),'true'); hp.append(repeat)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: font(r,size)
    geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(0); return t

def para(doc,text,bold_prefix=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.1
    if bold_prefix and text.startswith(bold_prefix): font(p.add_run(bold_prefix),10.5,bold=True,color=NAVY); text=text[len(bold_prefix):]
    font(p.add_run(text),10.5); return p

def heading(doc,text,level=1): return doc.add_paragraph(text,style=f'Heading {level}')

def callout(doc,label,text,color=GREEN):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8); p.paragraph_format.left_indent=Inches(.08); p.paragraph_format.right_indent=Inches(.08)
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F4F6F9'); p._p.get_or_add_pPr().append(shd)
    font(p.add_run(label+': '),10.5,bold=True,color=color); font(p.add_run(text),10.5)

def page_field(p):
    f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),'PAGE'); p._p.append(f)

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.49)
normal=doc.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,color,bef,aft in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',11.5,DARK,8,4)]:
    st=doc.styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=color; st.paragraph_format.space_before=Pt(bef); st.paragraph_format.space_after=Pt(aft); st.paragraph_format.keep_with_next=True
h=sec.header.paragraphs[0]; font(h.add_run('ANFA Research | Step 6 Native-K3s Actuation Delay'),9,color=GRAY)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(f.add_run('Page '),9,color=GRAY); page_field(f)

for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('CAPACITY ACTUATION\nDELAY REPORT'),28,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('ANFA Research Project - Step 6'),15,color=DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Native K3s cloud calibration on Microsoft Azure'),11,italic=True,color=GRAY)
for _ in range(4): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('30 formal cached-image trials | 6 August 2026\nFinal forecast horizon: 6 seconds'),11,bold=True,color=DARK)
doc.add_page_break()

heading(doc,'Executive result')
para(doc,'Step 6 measured the delay from a synthetic forecast and scaling decision to real serving capacity in the final three-node Azure K3s environment. Thirty formal trials covered scale-ups from one Pod to two, three, and four Pods. Every trial was valid, every expected new Pod was observed, all service probes succeeded, and no new Pod restarted.')
table(doc,['Decision','Cloud result','Basis'],[
    ('Forecast horizon','6 seconds',f'ceil({slowest_p95:.3f} s slowest P95 service + {uncertainty:.3f} s uncertainty + 2 s margin)'),
    ('Operational metric','Effective first service','Requires every requested new Pod to be Ready and observed serving /work'),
    ('Image policy','Pre-pull immutable image','Controls registry/network variability in the main experiment'),
    ('Applicability','Azure K3s environment','Supersedes the local kind timing constant for the cloud experiment')],[1900,2500,4960],8.7)
callout(doc,'Final decision','Use a 6-second forecast horizon for the Azure K3s main experiment with the benchmark image pre-pulled and verified on every worker.')

heading(doc,'1. Relationship to Steps 4 and 5')
para(doc,'Step 4 built the deterministic, stateless Go benchmark application. Step 5 established the empirical cloud capacity lookup C1=30, C2=40, C3=55, and C4=65 RPS. Step 6 determines how early the chosen replica count must be requested. Together, Step 5 answers “how many Pods?” and Step 6 answers “how early?”')
callout(doc,'Controller handoff','The controller converts forecast workload to replicas using the Step 5 lookup, then requests that capacity at least 6 seconds before the forecast traffic arrival.',DARK)

heading(doc,'2. Cloud environment and frozen inputs')
table(doc,['Component','Frozen configuration'],[
    ('Cloud','Microsoft Azure for Students; Malaysia West'),('Control plane','anfa-server, Standard_DS2_v2_Promo, Ubuntu 24.04'),('Workers','anfa-worker1 and anfa-worker2, Standard_D2_v2_Promo'),('Kubernetes','K3s v1.36.1+k3s1; one control plane and two workers'),('Application','Go benchmark v0.1.0; 50,000 SHA-256 iterations per /work request'),('Image','anfa/benchmark-app@sha256:0fd880c...61cee; linux/amd64'),('Resources','500m CPU and 128Mi memory request=limit'),('Probe path','/livez, /readyz, /work; NodePort 30080'),('Main image treatment','Image imported into K3s containerd and verified on every node; IfNotPresent'),('Monitoring','Prometheus Operator stack; separate from delay runner')],[2300,7060],8.7)

heading(doc,'3. Timestamp and delay model')
table(doc,['Event','Authority'],[
    ('Forecast available / controller decision','Native runner UTC on anfa-server'),('Scale sent / API acknowledged','Native runner UTC around k3s kubectl scale'),('Pod created','metadata.creationTimestamp'),('Scheduled','PodScheduled condition'),('Container started','containerStatuses.state.running.startedAt'),('Ready','Ready condition plus high-resolution runner observation'),('Application ready','X-Benchmark-Ready-At response header'),('First request served','Independent server-side 10 RPS probe and Pod UID header')],[4100,5260],8.7)
para(doc,'Kubernetes lifecycle timestamps have whole-second resolution, so zero or slightly negative raw component differences can occur inside the same second. Raw values are preserved for ordering and diagnosis. High-resolution server-side observation times are used for creation, readiness, and effective-serving totals.')

heading(doc,'4. Experimental method')
table(doc,['Factor','Method'],[
    ('Baseline','Exactly one Ready benchmark Pod'),('Treatments','1->2, 1->3, and 1->4 replicas'),('Repetitions','10 valid trials per treatment; 30 total'),('Order','Rotated 2/3/4, 3/4/2, and 4/2/3 to reduce time bias'),('Recovery','15 seconds before each scale action'),('Lifecycle polling','100 ms target using the native server-side runner'),('Service probing','Independent /work probe at 10 RPS from anfa-server'),('Completion','Every requested new Pod Ready and observed serving'),('Reset','Deployment automatically returned to one Pod after every trial'),('Evidence','Summary JSON, per-Pod CSV, service-probe CSV, Pods, Deployment, and Events')],[2200,7160],8.7)

heading(doc,'5. Instrumentation development and corrections')
table(doc,['Issue found','Correction and evidence treatment'],[
    ('Laptop-to-cloud latency dominated early load tests','Moved load and service probes inside Azure; rejected internet-path results'),('Inline Bash quoting corrupted timestamps','Marked the pilot invalid; installed a dedicated Python probe'),('Lifecycle polling reduced service-probe frequency','Separated probing into an independent thread; v2 pilot achieved 107 ms maximum gap'),('PowerShell progress collection became scalar','Preserved completed evidence, repaired resumable array handling, and recovered the missing progress row'),('Laptop restarted during campaign','Remote activity checked, one-Pod baseline restored, and campaign resumed from saved evidence'),('Whole-second API timestamps produced zero/negative raw fields','Retained raw diagnostics and used high-resolution observed totals')],[3000,6360],8.5)
callout(doc,'Evidence rule','Pilots and interrupted directories are preserved but excluded. Only 30 valid formal summaries contribute to the reported distributions.',AMBER)

heading(doc,'6. Campaign integrity')
table(doc,['Check','Result'],[
    ('Formal trials','30/30 valid'),('Repetitions','10 per replica target'),('Expected new Pods','60'),('Per-Pod evidence rows','60'),('New-Pod restarts','0'),('Distinct image IDs','1'),('Failed service probes','0'),('Maximum observed probe gap',f'{uncertainty*1000:.3f} ms'),('Final environment state','All three Azure VMs deallocated; resources and evidence preserved')],[5200,4160],9)

heading(doc,'7. Delay distributions')
table(doc,['Scale','Metric','n','Median','P90','P95','Maximum'],[(f'1->{n}',m,str(k),f'{med:.3f}',f'{p90:.3f}',f'{p95:.3f}',f'{mx:.3f}') for n,m,k,med,p90,p95,mx in dist],[1100,2100,650,1350,1350,1350,1460],8.3)
para(doc,'The 1->2 treatment produced the highest effective-serving P95 at 3.731 seconds. Larger increments were not slower in this environment because Pods were created in parallel and the Service probe sampled multiple backends rapidly. The horizon therefore uses the slowest observed treatment rather than assuming that the largest increment must dominate.')

heading(doc,'8. Forecast-horizon calculation')
callout(doc,'Predeclared rule','H = ceil(slowest treatment P95 effective-serving delay + measurement uncertainty + max(2 seconds, 20% of P95)).',DARK)
table(doc,['Term','Value'],[
    ('Slowest effective-serving P95',f'{slowest_p95:.3f} s (1->2)'),('Maximum probe-gap uncertainty',f'{uncertainty:.3f} s'),('20% proportional margin',f'{0.2*slowest_p95:.3f} s'),('Minimum margin','2.000 s'),('Selected margin',f'{margin:.3f} s'),('Unrounded total',f'{slowest_p95+uncertainty+margin:.3f} s'),('Final horizon',f'{horizon} seconds')],[3900,5460],9.5)
callout(doc,'Selected horizon',f'{horizon} seconds. This covers measured tail service delay, observed probe uncertainty, and the predeclared minimum safety margin.')

heading(doc,'9. Image-cache decision')
para(doc,'The formal cloud campaign used the immutable image pre-pulled into K3s containerd on all three nodes. All 60 new Pods used the same image ID. This treatment deliberately isolates Kubernetes scheduling, startup, readiness, EndpointSlice propagation, and Service participation from registry download variability.')
para(doc,'A separate cloud cold-pull campaign was not pooled with these results because no controlled immutable remote registry path was established. Earlier local Step 6 evidence showed that registry repull increased delay, but it is development evidence rather than a cloud estimate. The defensible main-experiment decision is therefore to pre-pull and verify the immutable image on every worker and retain IfNotPresent.')

heading(doc,'10. Findings')
table(doc,['Finding','Interpretation'],[
    ('API actuation is small','Pilot API acknowledgement was approximately 0.14-0.15 seconds.'),('Readiness is fast','Treatment P95 readiness ranged from 1.847 to 2.418 seconds.'),('Serving follows readiness','Treatment P95 first service ranged from 3.227 to 3.731 seconds.'),('Scale-up is parallel','The largest increment did not have the largest tail delay.'),('Probe quality is bounded','All probes succeeded; maximum observed gap was 164 ms.'),('Six seconds is conservative','It exceeds the slowest P95 by measurement uncertainty plus a two-second margin.')],[2900,6460],8.7)

heading(doc,'11. Limitations')
table(doc,['Limitation','Consequence'],[
    ('Small three-VM research cluster','Results apply to this Azure topology and VM family.'),('Two worker nodes','Four application replicas imply two Pods per worker, not four physical workers.'),('Whole-second lifecycle fields','Raw scheduling/startup subcomponents have limited sub-second precision.'),('Synthetic decision marker','Decision delay measures harness overhead, not the future controller implementation.'),('Service sampling','First service is an observed upper bound governed by probe timing and load balancing.'),('Cached-image main treatment','Registry outage and remote-pull variability are intentionally excluded.'),('Azure student SKUs','Promo availability and performance may differ from production VM families.')],[2700,6660],8.5)

heading(doc,'12. Reproducibility artifacts')
table(doc,['Artifact','Purpose'],[
    ('scripts/step6-cloud-runner.py','Single native-K3s trial; scaling, lifecycle polling, independent service probing'),('scripts/run-step6-cloud-campaign.ps1','Rotated, resumable 30-trial campaign and evidence download'),('step6/cloud/campaign-progress.csv','Accepted-trial ledger'),('step6/cloud/runs/cloud-cached-*','Immutable formal per-trial evidence'),('service-probe.csv','Every timestamped backend observation'),('per-pod.csv','Raw and high-resolution lifecycle fields'),('trial-summary.json','Trial-level readiness, service, probe, and validity metrics'),('Local Step 6 report','Separate development evidence; not pooled into the cloud distribution')],[3300,6060],8.5)

heading(doc,'13. Completion assessment')
table(doc,['Step 6 requirement','Status'],[
    ('Forecast, decision, scale, creation, scheduling, startup, readiness, and first-service timestamps','COMPLETE'),('Controlled scale-ups for different increments','COMPLETE'),('Repeated measurement and median/P90/P95/maximum','COMPLETE - 10 each'),('Image cached behavior and pre-pull decision','COMPLETE'),('Forecast horizon tied to measured serving delay','COMPLETE - 6 s'),('Variability and measurement limitations documented','COMPLETE'),('Capacity Actuation Delay Report','COMPLETE'),('Environment stopped without deleting evidence','COMPLETE')],[6900,2460],8.7)
callout(doc,'Official Step 6 result','Step 6 is complete for the Azure K3s research environment. Use H=6 seconds with the Step 5 empirical capacity lookup and the immutable image pre-pulled on every worker.',GREEN)

heading(doc,'Appendix A. Formal trial ledger')
ledger=sorted(valid,key=lambda r:(int(r['repetition']),int(r['target_replicas'])))
table(doc,['Rep','Scale','Ready (s)','Service (s)','Probe gap (ms)'],[(r['repetition'],f"1->{r['target_replicas']}",f"{r['readiness_delay_s']:.3f}",f"{r['effective_serving_delay_s']:.3f}",f"{r['maximum_probe_gap_ms']:.1f}") for r in ledger],[850,1500,2100,2300,2610],8.1)

heading(doc,'Appendix B. Final controller constants')
table(doc,['Input','Final cloud value'],[
    ('Step 5 empirical capacity','C1=30, C2=40, C3=55, C4=65 RPS'),('Step 6 forecast horizon','H=6 seconds'),('Replica rule','Use the smallest validated Pod count whose empirical capacity covers forecast workload'),('Timing rule','Request replicas at least 6 seconds before forecast workload arrival'),('Image rule','Pre-pull and verify the immutable benchmark image on every worker'),('Validated workload range','Up to 65 RPS and four benchmark Pods')],[3100,6260],9)

OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
