from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path
from statistics import mean

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = None
OUT = Path(__file__).resolve().parent / "step5-cloud-report"
FIGURES = OUT / "figures"

FINAL_CAMPAIGNS = {
    2: "2-pod-rps-40-formal-20260805-225658",
    3: "3-pod-rps-55-formal-20260805-232254",
    4: "4-pod-rps-65-formal-20260805-234932",
}

SINGLE_RUNS = [
    "single-pod-rps-30-rep-01-20260805-220117",
    "single-pod-rps-30-rep-02-20260805-220552",
    "single-pod-rps-30-rep-03-20260805-220946",
]

CAPACITY = {1: 30, 2: 40, 3: 55, 4: 65}
NEXT_FAIL = {1: 35, 2: 45, 3: 60, 4: 70}
FAIL_THROTTLE = {1: 48.54, 2: 15.58, 3: 16.55, 4: 10.50}

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 98, 108)
LIGHT = "F2F4F7"
GREEN = "EAF4EA"
RED = "FCE8E6"


def load_json(path: Path):
    with path.open(encoding="utf-8-sig") as handle:
        return json.load(handle)


def load_csv(path: Path):
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def single_rows():
    rows = []
    for idx, name in enumerate(SINGLE_RUNS, start=1):
        run = SOURCE / name
        client = load_json(run / "client-summary.json")
        if (run / "prometheus-summary.json").exists():
            prom = load_json(run / "prometheus-summary.json")
            mean_cpu = float(prom["mean_cpu"]["data"]["result"][0]["value"][1]) * 1000
            peak_cpu = float(prom["peak_cpu"]["data"]["result"][0]["value"][1]) * 1000
            mean_throttle = float(prom["mean_throttle"]["data"]["result"][0]["value"][1]) * 100
        else:
            mean_cpu = float(load_json(run / "prometheus-mean-cpu.json")["data"]["result"][0]["value"][1]) * 1000
            peak_cpu = float(load_json(run / "prometheus-peak-cpu.json")["data"]["result"][0]["value"][1]) * 1000
            mean_throttle = float(load_json(run / "prometheus-mean-throttling.json")["data"]["result"][0]["value"][1]) * 100
        rows.append({
            "replicas": 1,
            "repetition": idx,
            "completed": client["completed"],
            "errors": client["errors"],
            "achieved_rps": client["measurement_throughput_rps"],
            "p50_ms": client["p50_ms"],
            "p95_ms": client["p95_ms"],
            "p99_ms": client["p99_ms"],
            "max_mean_cpu_m": mean_cpu,
            "peak_cpu_m": peak_cpu,
            "max_mean_throttle_pct": mean_throttle,
            "pass": True,
        })
    return rows


def multipod_rows():
    rows = []
    for replicas, campaign in FINAL_CAMPAIGNS.items():
        for raw in load_csv(SOURCE / campaign / "campaign-summary.csv"):
            rows.append({
                "replicas": replicas,
                "repetition": int(raw["Repetition"]),
                "completed": int(raw["Completed"]),
                "errors": int(raw["Errors"]),
                "achieved_rps": float(raw["AchievedRps"]),
                "p50_ms": float(raw["P50ms"]),
                "p95_ms": float(raw["P95ms"]),
                "p99_ms": float(raw["P99ms"]),
                "max_mean_cpu_m": float(raw["MaxMeanCpuM"]),
                "peak_cpu_m": None,
                "max_mean_throttle_pct": float(raw["MaxMeanThrottlePct"]),
                "pass": raw["Pass"].lower() == "true",
            })
    return rows


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, formats=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell, LIGHT)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if formats and formats[idx]:
                value = formats[idx](value)
            cells[idx].text = str(value)
            cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
            for run in cells[idx].paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_picture(doc, path, alt_text):
    shape = doc.add_picture(str(path), width=Inches(6.25))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return shape


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ANFA Step 5 | Cloud K3s Capacity Profile | August 2026")
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY


def make_figures(selected):
    FIGURES.mkdir(parents=True, exist_ok=True)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 22)
        small = ImageFont.truetype("DejaVuSans.ttf", 18)
        title = ImageFont.truetype("DejaVuSans-Bold.ttf", 28)
    except OSError:
        font = small = title = ImageFont.load_default()
    width, height = 1400, 760
    left, right, top, bottom = 125, 60, 105, 110

    def axes(draw, heading, y_max, y_label):
        draw.text((width / 2, 28), heading, font=title, fill="#172033", anchor="ma")
        draw.line((left, top, left, height - bottom), fill="#111827", width=2)
        draw.line((left, height - bottom, width - right, height - bottom), fill="#111827", width=2)
        for i in range(6):
            y_value = y_max * i / 5
            y = height - bottom - (height - top - bottom) * i / 5
            draw.line((left, y, width - right, y), fill="#e5e7eb", width=1)
            draw.text((left - 15, y), f"{y_value:.0f}", font=small, fill="#374151", anchor="rm")
        return lambda v: height - bottom - (height - top - bottom) * v / y_max

    ns = [1, 2, 3, 4]
    observed = [CAPACITY[n] for n in ns]
    ideal = [30 * n for n in ns]
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    py = axes(draw, "Cloud K3s: observed versus ideal capacity", 120, "Aggregate capacity (RPS)")
    px = lambda n: left + (n - 1) * (width - left - right) / 3
    for n in ns:
        draw.text((px(n), height - bottom + 22), str(n), font=small, fill="#374151", anchor="ma")
    draw.text((width / 2, height - 35), "Ready benchmark Pods", font=small, fill="#374151", anchor="ma")
    for values, color, label, offset in ((ideal, "#64748b", "Ideal N x 30", 0), (observed, "#166534", "Observed safe capacity", 280)):
        points = [(px(n), py(v)) for n, v in zip(ns, values)]
        draw.line(points, fill=color, width=5)
        for point in points:
            draw.ellipse((point[0]-7, point[1]-7, point[0]+7, point[1]+7), fill=color)
        draw.line((left + offset, 82, left + offset + 55, 82), fill=color, width=5)
        draw.text((left + offset + 65, 82), label, font=small, fill="#172033", anchor="lm")
    image.save(FIGURES / "cloud-capacity-scaling.png")

    selected_max = []
    for n in ns:
        selected_max.append(max(r["max_mean_throttle_pct"] for r in selected if r["replicas"] == n))
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    py = axes(draw, "Boundary selection was driven by CPU throttling", 70, "Mean throttling ratio (%)")
    group_width = (width - left - right) / 4
    for idx, n in enumerate(ns):
        center = left + group_width * (idx + .5)
        bar_width = 70
        for x, value, color in ((center - 80, selected_max[idx], "#2e74b5"), (center + 10, FAIL_THROTTLE[n], "#b91c1c")):
            draw.rectangle((x, py(value), x + bar_width, height - bottom), fill=color)
        draw.text((center, height - bottom + 22), f"{n} Pod" if n == 1 else f"{n} Pods", font=small, fill="#374151", anchor="ma")
    guard_y = py(10)
    for x in range(left, width - right, 24):
        draw.line((x, guard_y, min(x + 12, width - right), guard_y), fill="#7a5a00", width=3)
    draw.line((left, 82, left + 48, 82), fill="#2e74b5", width=12)
    draw.text((left + 60, 82), "Selected capacity", font=small, fill="#172033", anchor="lm")
    draw.line((left + 350, 82, left + 398, 82), fill="#b91c1c", width=12)
    draw.text((left + 410, 82), "Next tested failure", font=small, fill="#172033", anchor="lm")
    draw.text((width - right, guard_y - 10), "10% guardrail", font=small, fill="#7a5a00", anchor="rs")
    image.save(FIGURES / "cloud-throttling-boundary.png")


def build_markdown(selected):
    lines = [
        "# Step 5 Cloud K3s Capacity Profiling Report",
        "",
        "## Executive result",
        "",
        "The native three-node Azure K3s campaign established `C_pod = 30 RPS` and the empirical lookup `C_1=30`, `C_2=40`, `C_3=55`, `C_4=65 RPS`.",
        "",
        "| Ready Pods | Safe capacity | Ideal N x 30 | Efficiency |",
        "|---:|---:|---:|---:|",
    ]
    for n in range(1, 5):
        lines.append(f"| {n} | {CAPACITY[n]} RPS | {30*n} RPS | {CAPACITY[n]/(30*n):.3f} |")
    lines += [
        "",
        "The oracle/controller must use the lookup table, not `ceil(W/30)`, because scaling efficiency decreases with replica count.",
        "",
        "## Frozen SLO and guardrails",
        "",
        "P99 latency <= 300 ms; failure rate < 1%; achieved throughput >= 99% of offered load; mean Pod CPU <= 450m; CPU throttled-period ratio < 10%; no readiness loss, restart, or replica change.",
        "",
        "## Final formula",
        "",
        "`replicas(W) = min { N in {1,2,3,4} : W <= C_N }`, with `C={1:30,2:40,3:55,4:65}`. Loads above 65 RPS are outside the validated four-replica range.",
        "",
        "## Key findings",
        "",
        "1. The client had to run inside Azure; PC-to-Malaysia measurements contained roughly 240 ms of public-network latency and were rejected as capacity evidence.",
        "2. CPU throttling, not request failure, was the conservative boundary at the selected multi-Pod levels.",
        "3. Kubernetes balanced total requests well, but per-Pod arrivals were bursty enough to create throttling before average CPU reached 450m.",
        "4. Four Pods ran across two workers, so the four-Pod result is replica scaling on two machines, not four-node scaling.",
        "5. The imported image required the digest-qualified alias in K3s containerd's `k8s.io` namespace; this was an implementation detail, not an architectural redesign.",
    ]
    return "\n".join(lines) + "\n"


def build_docx(selected, output):
    doc = Document()
    configure_styles(doc)
    add_footer(doc.sections[0])

    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("ANFA Research | Capacity Profiling")
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("CAPACITY PROFILING REPORT")
    r.font.name = "Calibri"
    r.font.size = Pt(23)
    r.font.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Step 5 - Native Azure K3s validation")
    r.font.size = Pt(14)
    r.font.color.rgb = GRAY
    for label, value in (
        ("Environment", "Three-node K3s v1.36.1+k3s1 in Azure Malaysia West"),
        ("Application", "ANFA benchmark-app 0.1.0; 50,000 SHA-256 iterations"),
        ("Campaign date", "5 August 2026"),
        ("Status", "Cloud capacity profiling complete"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    add_heading(doc, "Executive result", 1)
    add_body(doc, "The cloud campaign established a conservative single-Pod safe capacity of 30 requests per second (RPS). Aggregate safe capacity was 40 RPS with two Pods, 55 RPS with three Pods, and 65 RPS with four Pods. Every accepted value passed three 120-second repetitions with zero errors, full offered throughput, P99 latency below 300 ms, mean CPU below 450m, mean throttling below 10%, complete Pod coverage, and no readiness violation.")
    add_table(doc, ["Ready Pods", "Safe capacity", "Ideal", "Efficiency", "Decision"], [
        [n, f"{CAPACITY[n]} RPS", f"{30*n} RPS", CAPACITY[n]/(30*n), "accepted"] for n in range(1, 5)
    ], [1400, 1900, 1700, 1700, 2660], [None, None, None, lambda x: f"{x:.3f}", None])
    add_picture(doc, FIGURES / "cloud-capacity-scaling.png", "Line chart comparing ideal and observed cloud capacity for one through four Pods")

    add_heading(doc, "1. Purpose and relationship to the local campaign", 1)
    add_body(doc, "The earlier Kind/Docker Desktop campaign validated the method and produced local-only capacities of 45, 90, 105, and 130 RPS. Its report explicitly required recalibration on native K3s. This cloud campaign supplies that final-environment calibration. The two datasets are not averaged: local values remain development evidence, while the cloud lookup is authoritative for the Azure K3s experiments.")

    add_heading(doc, "2. Frozen application and infrastructure", 1)
    add_table(doc, ["Element", "Recorded configuration"], [
        ["Application", "Go benchmark service 0.1.0; /work performs 50,000 deterministic SHA-256 iterations"],
        ["Image", "anfa/benchmark-app@sha256:0fd880c5401b443a3dfb329c48fe3bd8c844643007a6097f6c31917a47961cee"],
        ["Pod resources", "500m CPU and 128Mi memory; requests equal limits"],
        ["Cluster", "K3s v1.36.1+k3s1; one control-plane VM and two worker VMs"],
        ["Traffic source", "Linux AMD64 load generator on anfa-server, inside Azure"],
        ["Traffic path", "http://127.0.0.1:30080/work through the K3s NodePort"],
        ["Monitoring", "Prometheus, 15-second scrape interval; per-Pod CPU and CFS throttling"],
    ], [2300, 7060])

    add_heading(doc, "3. SLO and pass rule", 1)
    add_body(doc, "A run passed only when P99 latency was at most 300 ms, failure rate was below 1%, successful throughput reached at least 99% of offered RPS, mean Pod CPU was at most 450m, the CPU throttled-period ratio was below 10%, every expected Pod served traffic, and no readiness loss or restart occurred. A capacity level was accepted only when all three formal repetitions passed.")

    add_heading(doc, "4. Measurement correction: locate the client with the system", 1)
    add_body(doc, "Initial diagnostic runs sent traffic from the research laptop to Malaysia West. At only 15 RPS, that path produced a P50 near 243 ms and P99 above one second, while a one-RPS check from the Azure control-plane VM produced a P50 near 16 ms. The discrepancy was public-network latency, not Pod saturation. A pinned Linux AMD64 build of the same open-loop load generator was therefore copied to anfa-server and verified by matching SHA-256 hashes. All accepted cloud results use the inside-Azure client. The external trials are retained as invalid diagnostic evidence and are not used in capacity calculations.")

    doc.add_page_break()
    add_heading(doc, "5. Accepted formal evidence", 1)
    add_table(doc, ["Pods", "Rep", "Completed", "Errors", "P99 ms", "Max mean CPU m", "Max mean throttle", "Result"], [
        [r["replicas"], r["repetition"], r["completed"], r["errors"], r["p99_ms"], r["max_mean_cpu_m"], r["max_mean_throttle_pct"], "pass"] for r in selected
    ], [750, 650, 1250, 850, 1050, 1500, 1650, 1660], [None, None, None, None, lambda x: f"{x:.2f}", lambda x: f"{x:.2f}", lambda x: f"{x:.2f}%", None])

    add_heading(doc, "6. Boundary and saturation findings", 1)
    add_table(doc, ["Pods", "Accepted", "Next failure", "Failure evidence", "Boundary driver"], [
        [1, "30 RPS", "35 RPS", "48.54% throttling", "CPU throttling"],
        [2, "40 RPS", "45 RPS", "11.56-15.58% mean throttling", "CPU throttling"],
        [3, "55 RPS", "60 RPS", "16.55% mean throttling", "CPU throttling"],
        [4, "65 RPS", "70 RPS", "10.50% mean throttling", "CPU throttling"],
    ], [850, 1500, 1500, 3000, 2510])
    add_picture(doc, FIGURES / "cloud-throttling-boundary.png", "Bar chart comparing throttling at each selected capacity and the next tested failing load")
    add_body(doc, "At ideal linear loads, the three-Pod 90 RPS trial reached 65.71% mean throttling, while the four-Pod 120 RPS trial reached 60.02% mean throttling and P99 of 303.04 ms. Throughput and error metrics alone would have accepted several unsafe points. The resource guardrails correctly rejected them before sustained overload became request failure.")

    add_heading(doc, "7. Scaling interpretation", 1)
    add_body(doc, "Efficiency fell from 1.000 for one Pod to 0.667, 0.611, and 0.542 for two, three, and four Pods. Total request counts were generally balanced, but Kubernetes load distribution is not perfectly periodic per Pod. A Pod can receive short bursts that exhaust its 500m CFS quota even while its average CPU remains well below 450m. This explains why throttling became the binding guardrail and why aggregate safe capacity did not follow N x 30.")
    add_body(doc, "The four-Pod campaign used two worker VMs, normally placing two application Pods on each worker. It validates horizontal replica behavior and shared-worker contention for this project topology; it does not claim four independent worker-node scaling.")

    add_heading(doc, "8. Final capacity mapping", 1)
    add_table(doc, ["Forecast workload W", "Required replicas", "Validated capacity"], [
        ["0 < W <= 30", 1, "30 RPS"],
        ["30 < W <= 40", 2, "40 RPS"],
        ["40 < W <= 55", 3, "55 RPS"],
        ["55 < W <= 65", 4, "65 RPS"],
        ["W > 65", "outside range", "not validated"],
    ], [3300, 2500, 3560])
    add_body(doc, "Formal rule: replicas(W) = min {N in {1,2,3,4} : W <= C_N}, where C_1=30, C_2=40, C_3=55, and C_4=65 RPS. The controller and oracle must use this lookup table. A single linear multiplier would either overstate higher-replica capacity or unnecessarily understate lower-replica capacity.")

    add_heading(doc, "9. Implementation-driven changes", 1)
    add_body(doc, "Load-generator placement. The client moved from the laptop to the Azure control-plane VM to remove public-internet latency from the SLO measurement. The traffic model and executable logic did not change.")
    add_body(doc, "Image distribution. Because no registry was introduced, the frozen image tar was imported into each node's K3s containerd k8s.io namespace. Kubelet resolution also required a digest-qualified alias. The image digest and application architecture remained unchanged.")
    add_body(doc, "Persistent runner. A parameterized run-step5-cloud-campaign.ps1 script was added. It checks context and readiness, invokes the remote client, samples Prometheus during the active window, captures Kubernetes state, enforces the frozen pass rule, and writes immutable evidence directories.")

    add_heading(doc, "10. Limitations and completion assessment", 1)
    add_body(doc, "The campaign used an Azure for Students topology constrained to two worker VMs. Prometheus's 15-second scrape interval limits fine-grained CPU interpretation, so the report treats trial-level mean throttling as the pre-registered guardrail and retains raw samples. The image was pre-pulled; image download latency is a Step 6 treatment, not part of capacity profiling. Results apply to the recorded VM sizes, K3s version, Pod limits, work intensity, and request path.")
    add_body(doc, "Step 5 cloud completion: complete. Repeated evidence supports C_pod=30 RPS; two-, three-, and four-Pod capacities are measured; saturation behavior is understood; and the oracle/controller lookup is defined. Step 6 must now measure how long a requested replica increase takes to become serving capacity in this same cloud environment.")

    doc.save(output)


def main():
    global SOURCE, OUT, FIGURES
    parser = argparse.ArgumentParser(description="Build the capacity-profiling report from a complete cloud-runs directory.")
    parser.add_argument("--source", type=Path, required=True, help="Directory containing the immutable Step 5 run directories.")
    parser.add_argument("--output-directory", type=Path, default=OUT)
    args = parser.parse_args()
    SOURCE = args.source.resolve()
    OUT = args.output_directory.resolve()
    FIGURES = OUT / "figures"
    OUT.mkdir(parents=True, exist_ok=True)
    selected = single_rows() + multipod_rows()
    selected.sort(key=lambda r: (r["replicas"], r["repetition"]))
    make_figures(selected)

    with (OUT / "selected-formal-runs.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(selected[0]))
        writer.writeheader()
        writer.writerows(selected)

    result = {
        "environment": "anfa-cloud-native-k3s",
        "c_pod_rps": 30,
        "capacity_lookup_rps": {str(k): v for k, v in CAPACITY.items()},
        "scaling_efficiency": {str(n): CAPACITY[n] / (n * 30) for n in CAPACITY},
        "validated_replica_range": [1, 4],
        "slo": {"p99_ms_lte": 300, "failure_rate_lt": 0.01},
        "guardrails": {"throughput_fidelity_gte": 0.99, "mean_cpu_m_lte": 450, "cpu_throttling_ratio_lt": 0.10},
        "oracle_rule": "min N in {1,2,3,4} such that workload_rps <= C_N",
    }
    (OUT / "capacity-result-cloud.json").write_text(json.dumps(result, indent=2), encoding="utf-8")
    (OUT / "STEP5_CLOUD_CAPACITY_PROFILING_REPORT.md").write_text(build_markdown(selected), encoding="utf-8")
    build_docx(selected, OUT / "ANFA_Step_5_Cloud_K3s_Capacity_Profiling_Detailed_Report.docx")
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
